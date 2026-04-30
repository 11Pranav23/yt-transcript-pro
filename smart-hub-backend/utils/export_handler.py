"""
Export handlers for different transcript formats
"""
import os
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT

def generate_srt(text, start_time=0, chunk_duration=5):
    """
    Generate SRT subtitle format from text
    
    Args:
        text: Transcript text
        start_time: Starting time in seconds
        chunk_duration: Duration of each subtitle in seconds
    
    Returns: SRT formatted string
    """
    try:
        sentences = text.split('.')
        sentences = [s.strip() for s in sentences if s.strip()]
        
        srt_content = []
        current_time = start_time
        index = 1
        
        for sentence in sentences:
            if not sentence:
                continue
            
            start = format_time_srt(current_time)
            end = format_time_srt(current_time + chunk_duration)
            
            srt_content.append(f"{index}\n{start} --> {end}\n{sentence}\n")
            
            current_time += chunk_duration
            index += 1
        
        return '\n'.join(srt_content)
        
    except Exception as e:
        print(f"Error generating SRT: {str(e)}")
        return ""

def generate_vtt(text, start_time=0, chunk_duration=5):
    """
    Generate VTT subtitle format from text
    """
    try:
        sentences = text.split('.')
        sentences = [s.strip() for s in sentences if s.strip()]
        
        vtt_content = ['WEBVTT\n']
        current_time = start_time
        
        for sentence in sentences:
            if not sentence:
                continue
            
            start = format_time_vtt(current_time)
            end = format_time_vtt(current_time + chunk_duration)
            
            vtt_content.append(f"{start} --> {end}\n{sentence}\n")
            
            current_time += chunk_duration
        
        return '\n'.join(vtt_content)
        
    except Exception as e:
        print(f"Error generating VTT: {str(e)}")
        return ""

def generate_pdf(text, title="Transcript", filename="transcript"):
    """
    Generate PDF document from transcript
    """
    try:
        output_path = os.path.join(
            os.path.dirname(__file__), '..',
            'uploads',
            f"{filename}.pdf"
        )
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1f2937',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Header
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 0.3 * inch))
        
        # Body text
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=12
        )
        
        # Split text into paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
        
        # Build PDF
        doc.build(story)
        
        return output_path
        
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return None

def generate_docx(text, title="Transcript", filename="transcript"):
    """
    Generate DOCX Word document from transcript
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        output_path = os.path.join(
            os.path.dirname(__file__), '..',
            'uploads',
            f"{filename}.docx"
        )
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Create document
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph(title)
        title_para.style = 'Heading 1'
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add date
        date_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add content
        doc.add_paragraph()  # Blank line
        doc.add_paragraph(text)
        
        # Save
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return None

def format_time_srt(seconds):
    """Format time for SRT (HH:MM:SS,mmm)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"

def format_time_vtt(seconds):
    """Format time for VTT (HH:MM:SS.mmm)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"

def generate_markdown(text, title="Transcript"):
    """
    Generate Markdown format transcript
    """
    try:
        markdown_content = f"# {title}\n\n"
        markdown_content += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        markdown_content += "---\n\n"
        markdown_content += text
        
        return markdown_content
    except Exception as e:
        print(f"Error generating Markdown: {str(e)}")
        return text

def generate_json_export(transcript_data, metadata=None):
    """
    Generate JSON export with metadata
    """
    try:
        import json
        
        export_data = {
            'transcript': transcript_data,
            'metadata': metadata or {},
            'generated_at': datetime.now().isoformat(),
            'version': '1.0'
        }
        
        return json.dumps(export_data, indent=2)
    except Exception as e:
        print(f"Error generating JSON: {str(e)}")
        return ""
